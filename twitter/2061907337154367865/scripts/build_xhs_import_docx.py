from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "02_小红书发布正文.md"
ASSET_DIR = ROOT / "assets"
OUT_PATH = ROOT / "小红书导入版_图文交错.docx"


IMAGE_BEFORE_HEADING = {
    "# Codex 动态工作流：让 AI 为每个任务临时搭一套执行系统": [
        "01_封面_Codex动态工作流.png",
        "00_底稿完整内容总览.png",
    ],
    "## 先看几个典型用法": ["04_适合workflow的任务.png"],
    "## 动态工作流到底怎么运行": ["05_workflow运行机制.png"],
    "## 为什么需要动态工作流": ["03_复杂任务压垮单上下文.png"],
    "## 动态工作流和静态工作流有什么不同": ["07_动态vs静态workflow.png"],
    "### 模式一：Classify-and-act": ["08_先分类再行动.png"],
    "### 模式二：Fan-out-and-synthesize": ["09_拆开并行最后合并.png"],
    "### 模式三：Adversarial verification": ["10_生成者和验证者分开.png"],
    "### 模式四：Generate-and-filter": ["11_生成过滤.png"],
    "### 模式五：Tournament": ["12_tournament方案比赛.png"],
    "### 模式六：Loop until done": ["13_循环到真正完成.png"],
    "## 具体怎么用：6 个可以直接照抄的案例": ["16_怎么让Codex用好workflow.png"],
    "## 动态工作流适合哪些场景": ["14_workflow任务地图.png"],
    "## 什么时候不要使用动态工作流": ["15_不是每个任务都需要workflow.png"],
    "### 保存和分享 workflow": ["17_workflow沉淀成资产.png"],
    "## 最后总结": ["18_从AI助手到任务指挥台.png"],
}

IMAGE_AFTER_PARAGRAPH_CONTAINS = {
    "第三类是 goal drift": ["06_单agent三个问题.png"],
    "它不是先做一个万能模板": ["02_不是一个Prompt.png"],
}


def set_east_asia_font(run, font_name="PingFang SC"):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text_run(paragraph, text, bold=False, size=11, color=None):
    run = paragraph.add_run(text)
    set_east_asia_font(run)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_image(doc, filename):
    path = ASSET_DIR / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=8, after=4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(4.2))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(caption, before=0, after=10)
    stem = path.stem
    add_text_run(caption, stem, size=9, color="666666")


def add_heading(doc, text, level):
    if level == 1:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=12, after=8, line=1.15)
        add_text_run(p, text, bold=True, size=20, color="111111")
    elif level == 2:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=14, after=6, line=1.15)
        add_text_run(p, text, bold=True, size=16, color="111111")
    else:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=10, after=5, line=1.15)
        add_text_run(p, text, bold=True, size=13, color="333333")
    return p


def add_paragraph_markdown(doc, line):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=7, line=1.18)
    # Minimal inline code handling.
    parts = re.split(r"(`[^`]+`)", line)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = add_text_run(p, part[1:-1], size=10, color="333333")
            run.font.name = "Menlo"
        else:
            add_text_run(p, part, size=11, color="111111")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    set_paragraph_spacing(p, before=0, after=4, line=1.15)
    add_text_run(p, "• ", size=11, color="111111")
    add_text_run(p, text, size=11, color="111111")


def add_code_block(doc, lines):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.right_indent = Inches(0.1)
        set_paragraph_spacing(p, before=4 if i == 0 else 0, after=4 if i == len(lines) - 1 else 0, line=1.05)
        run = p.add_run(line)
        run.font.name = "Menlo"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(50, 50, 50)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "PingFang SC"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        if not line:
            continue

        if line in IMAGE_BEFORE_HEADING:
            for image in IMAGE_BEFORE_HEADING[line]:
                add_image(doc, image)

        if line.startswith("# "):
            add_heading(doc, line[2:], 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:], 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:], 3)
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            set_paragraph_spacing(p, before=0, after=4, line=1.15)
            add_text_run(p, line, size=11, color="111111")
        elif line.startswith("- "):
            add_bullet(doc, line[2:])
        else:
            add_paragraph_markdown(doc, line)

        for marker, images in IMAGE_AFTER_PARAGRAPH_CONTAINS.items():
            if marker in line:
                for image in images:
                    add_image(doc, image)

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_docx()
    print(OUT_PATH)
